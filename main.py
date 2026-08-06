import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from openai import OpenAI
import streamlit as str_app

# חיבור למנוע של OpenAI
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

# עיצוב ממשק האתר
str_app.title("תמלול הקלטות לקובץ Word")
str_app.write("העלה קובץ אודיו וקבל קובץ וורד מתומלל בעברית תקינה")

uploaded_file = str_app.file_uploader("בחר קובץ אודיו (mp3, wav, m4a)", type=["mp3", "wav", "m4a"])

if uploaded_file is not None:
    if str_app.button("התחל תמלול"):
        with str_app.spinner("ה-AI מתמלל את הקובץ שלך כעת... אנא המתן"):
            
            # תיקון השגיאה: שמירת הקובץ בשם אנגלי קבוע ללא תלות בשם המקורי בעברית
            temp_filename = "temp_audio_file.mp3"
            with open(temp_filename, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            try:
                # שליחה ל-Whisper
                with open(temp_filename, "rb") as audio_file:
                    transcription = client.audio.transcriptions.create(
                        model="whisper-1", 
                        file=audio_file,
                        language="he"
                    )
                
                text_result = transcription.text
                
                # יצירת קובץ Word מעוצב לעברית (RTL)
                doc = docx.Document()
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                pPr = paragraph._p.get_or_add_pPr()
                pPr.append(OxmlElement('w:bidi'))
                
                run = paragraph.add_run(text_result)
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(12)
                
                output_filename = "output_transcription.docx"
                doc.save(output_filename)
                
                # כפתור להורדת הקובץ המוכן
                with open(output_filename, "rb") as word_file:
                    str_app.download_button(
                        label="📥 הורד קובץ Word מוכן",
                        data=word_file,
                        file_name="תמלול_מוכן.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                str_app.success("התמלול הסתיים בהצלחה!")
                
                # ניקוי קבצים זמניים מהשרת
                if os.path.exists(temp_filename):
                    os.remove(temp_filename)
                if os.path.exists(output_filename):
                    os.remove(output_filename)
                    
            except Exception as e:
                str_app.error(f"שגיאה בתהליך: {e}")
   
